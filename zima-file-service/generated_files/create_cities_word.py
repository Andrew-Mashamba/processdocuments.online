#!/usr/bin/env python3

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add title
title = doc.add_heading('Five Famous Cities', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_heading('A Collection of Notable World Cities', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction paragraph
doc.add_paragraph('Here are five remarkable cities from different continents, each with its own unique character and attractions:')

# Add the five cities with descriptions
cities = [
    {
        'name': 'London, United Kingdom',
        'description': 'The capital of England, famous for Big Ben, Tower Bridge, and its rich royal history. A global financial center with incredible museums and parks.'
    },
    {
        'name': 'Barcelona, Spain',
        'description': 'Known for its stunning architecture by Antoni Gaudí, beautiful beaches, vibrant nightlife, and delicious tapas cuisine.'
    },
    {
        'name': 'Singapore',
        'description': 'A modern city-state in Southeast Asia, renowned for its clean streets, diverse food culture, and impressive skyline with Marina Bay Sands.'
    },
    {
        'name': 'Cape Town, South Africa',
        'description': 'Located at the southern tip of Africa with Table Mountain as its backdrop, famous for its wine regions and stunning coastal scenery.'
    },
    {
        'name': 'San Francisco, USA',
        'description': 'Home to the iconic Golden Gate Bridge, historic cable cars, and the tech hub of Silicon Valley. Known for its hilly terrain and foggy weather.'
    }
]

# Add each city as a numbered list item with description
for i, city in enumerate(cities, 1):
    # Add city name as a heading
    city_heading = doc.add_heading(f"{i}. {city['name']}", level=3)

    # Add description paragraph
    doc.add_paragraph(city['description'])

    # Add a small space between entries
    doc.add_paragraph("")

# Add closing paragraph
doc.add_paragraph('These five cities represent different cultures, architectural styles, and geographical regions, making them fascinating destinations for travelers and urban enthusiasts alike.')

# Save the document
output_path = 'generated_files/five_cities.docx'
doc.save(output_path)

print(f"Word document saved successfully: {output_path}")