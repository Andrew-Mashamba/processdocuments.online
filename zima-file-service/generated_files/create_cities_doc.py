#!/usr/bin/env python3
"""
Create a Word document with five beautiful cities around the world.
"""

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add title
title = doc.add_heading('Five Beautiful Cities Around the World', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction paragraph
intro = doc.add_paragraph()
intro.add_run('Here are five amazing cities from different continents, each offering unique experiences and attractions:')

# Add the five cities
cities = [
    {
        'name': 'Tokyo, Japan',
        'description': 'The bustling capital known for its blend of traditional culture and modern technology, featuring ancient temples alongside futuristic skyscrapers.'
    },
    {
        'name': 'Paris, France',
        'description': 'The romantic City of Light famous for its art, fashion, and iconic landmarks like the Eiffel Tower and Louvre Museum.'
    },
    {
        'name': 'New York City, USA',
        'description': 'The energetic metropolis that never sleeps, home to Broadway, Central Park, Times Square, and the Statue of Liberty.'
    },
    {
        'name': 'Sydney, Australia',
        'description': 'The harbor city renowned for its stunning Opera House, beautiful beaches, and the famous Harbor Bridge.'
    },
    {
        'name': 'Cairo, Egypt',
        'description': 'The ancient city rich in history, home to the Great Pyramids, the Sphinx, and countless archaeological treasures.'
    }
]

# Add each city
for i, city in enumerate(cities, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. {city['name']} - ").bold = True
    p.add_run(city['description'])

# Add conclusion paragraph
conclusion = doc.add_paragraph()
conclusion.add_run('Each of these cities represents a unique blend of culture, history, and modern attractions that make them must-visit destinations for travelers from around the world.')

# Save the document
output_path = "/Volumes/DATA/QWEN/zima-file-service/generated_files/cities.docx"
doc.save(output_path)

print(f"Word document created successfully: {output_path}")
print(f"Document contains {len(cities)} cities with descriptions.")